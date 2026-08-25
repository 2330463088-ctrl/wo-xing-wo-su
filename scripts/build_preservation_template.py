from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


PARAGRAPH_TEXT = {
    0: "诉前财产保全申请书",
    1: "申请人：{{APPLICANT_INFO}}",
    2: "被申请人：{{RESPONDENT_INFO}}",
    3: "请求事项：",
    4: "请求依法冻结被申请人银行存款{{PRESERVATION_AMOUNT}}元人民币，或查封、扣押其同等价值的财产。",
    5: "事实和理由： ",
    6: (
        "申请人与被申请人因民间借贷纠纷一案，申请人拟向{{COURT_NAME}}提起诉讼。"
        "为防止被申请人转移财产，保证生效裁判的顺利执行，保护申请人的合法权益，"
        "根据民事诉讼法的相关规定，现依法申请诉前财产保全。如保全不当，"
        "申请人自愿承担相应的法律责任。请批准。"
    ),
    7: "此致",
    8: "{{COURT_NAME}}",
    9: "申请人：{{APPLICANT_NAME}}        ",
    10: "{{YEAR}}年{{MONTH}}月{{DAY}}日",
}


def replace_paragraph_text(paragraph, text: str) -> None:
    """Replace text while retaining the paragraph and first run formatting."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()

    document = Document(args.source)
    if len(document.paragraphs) < 11:
        raise ValueError("参考文档段落不足，无法安全生成模板")
    for index, text in PARAGRAPH_TEXT.items():
        replace_paragraph_text(document.paragraphs[index], text)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)


if __name__ == "__main__":
    main()
